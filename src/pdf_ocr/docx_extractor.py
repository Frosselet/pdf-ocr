"""Word (DOCX/DOC) table extraction.

Extracts structured tables from Word documents.
- DOCX (Office 2007+): Uses python-docx
- DOC (Office 97-2003): Converts via LibreOffice if available, or raises error

DOCX tables have explicit grid structure, but require handling of:
- Horizontal merges (gridSpan)
- Vertical merges (vMerge)
- Nested tables (treated as text content)

Semantic heuristics from heuristics.py apply for header detection
and type pattern analysis.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING

from pdf_ocr.heuristics import (
    StructuredTable,
    TableMetadata,
    build_column_names_from_headers,
    estimate_header_rows,
    normalize_grid,
    split_headers_from_data,
)

if TYPE_CHECKING:
    from docx.document import Document
    from docx.table import Table, _Cell


# ---------------------------------------------------------------------------
# Visual element structures
# ---------------------------------------------------------------------------


@dataclass
class DocxCellStyle:
    """Visual style information for a table cell."""

    shading_color: tuple[float, float, float] | None = None  # RGB 0-1
    border_top: bool = False
    border_bottom: bool = False
    border_left: bool = False
    border_right: bool = False
    is_bold: bool = False
    is_italic: bool = False
    font_size: float | None = None


@dataclass
class DocxVisualInfo:
    """Visual information for the entire table."""

    header_fill_rows: list[int] = field(default_factory=list)
    has_borders: bool = False


@dataclass
class DocxTable:
    """A table extracted from a Word document with visual info."""

    grid: list[list[str]]
    styles: list[list[DocxCellStyle]] | None = None
    visual: DocxVisualInfo | None = None
    table_index: int = 0


# ---------------------------------------------------------------------------
# Color utilities
# ---------------------------------------------------------------------------


def _parse_hex_color(hex_color: str | None) -> tuple[float, float, float] | None:
    """Parse hex color string to RGB tuple (0-1 range).

    Handles formats: RRGGBB, #RRGGBB, auto, None
    """
    if not hex_color or hex_color.lower() in ("auto", "none"):
        return None

    # Strip # if present
    hex_color = hex_color.lstrip("#")

    if len(hex_color) != 6:
        return None

    try:
        r = int(hex_color[0:2], 16) / 255.0
        g = int(hex_color[2:4], 16) / 255.0
        b = int(hex_color[4:6], 16) / 255.0
        # Skip pure white (typically means no fill)
        if (r, g, b) == (1.0, 1.0, 1.0):
            return None
        return (r, g, b)
    except ValueError:
        return None


def _extract_cell_style(cell: "_Cell") -> DocxCellStyle:
    """Extract visual style information from a cell."""
    style = DocxCellStyle()

    # Shading (fill color)
    try:
        tc = cell._tc
        tc_pr = tc.tcPr
        if tc_pr is not None:
            shading = tc_pr.shd
            if shading is not None:
                fill = shading.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                if fill:
                    style.shading_color = _parse_hex_color(fill)
    except Exception:
        pass

    # Check paragraphs for font info
    try:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.bold:
                    style.is_bold = True
                if run.italic:
                    style.is_italic = True
                if run.font.size:
                    style.font_size = run.font.size.pt
    except Exception:
        pass

    return style


# ---------------------------------------------------------------------------
# Merged cell handling
# ---------------------------------------------------------------------------


def _get_cell_text(cell: "_Cell") -> str:
    """Get text content from a cell, handling nested tables."""
    # Collect text from all paragraphs
    text_parts: list[str] = []
    for para in cell.paragraphs:
        # split()+join collapses internal runs of whitespace to single spaces
        text = " ".join(para.text.split())
        if text:
            text_parts.append(text)

    # Note: Nested tables are ignored (their text is in paragraphs)
    return " ".join(text_parts)


def _build_grid_from_table(table: "Table") -> tuple[list[list[str]], list[list[DocxCellStyle]]]:
    """Build a grid from a Word table, handling merged cells.

    Word tables use:
    - gridSpan for horizontal merges (cell spans multiple columns)
    - vMerge for vertical merges (cell continues from above)

    Returns (grid, styles) where grid is list of rows, each row is list of strings.
    """
    # First pass: determine grid dimensions
    num_rows = len(table.rows)
    if num_rows == 0:
        return [], []

    # Count columns from first row (gridSpan-aware)
    # python-docx returns duplicate _tc references for horizontally merged cells
    # (a cell with gridSpan=2 appears twice in row.cells, both pointing to the
    # same XML element). We must deduplicate by tracking seen _tc ids.
    num_cols = 0
    seen_tcs: set[int] = set()
    for cell in table.rows[0].cells:
        tc_id = id(cell._tc)
        if tc_id in seen_tcs:
            continue
        seen_tcs.add(tc_id)
        try:
            tc = cell._tc
            tc_pr = tc.tcPr
            grid_span = tc_pr.gridSpan if tc_pr is not None else None
            span = int(grid_span.val) if grid_span is not None else 1
            num_cols += span
        except Exception:
            num_cols += 1

    if num_cols == 0:
        return [], []

    # Initialize grid
    grid: list[list[str]] = [[""] * num_cols for _ in range(num_rows)]
    styles: list[list[DocxCellStyle]] = [[DocxCellStyle() for _ in range(num_cols)] for _ in range(num_rows)]

    # Track vertical merge state per column
    # vmerge_continue[col] = (row, text, style) of the cell that started the merge
    vmerge_start: dict[int, tuple[int, str, DocxCellStyle]] = {}

    for row_idx, row in enumerate(table.rows):
        col_idx = 0
        seen_tcs_row: set[int] = set()

        for cell in row.cells:
            # Deduplicate: python-docx returns the same _tc for merged cells
            tc_id = id(cell._tc)
            if tc_id in seen_tcs_row:
                continue
            seen_tcs_row.add(tc_id)

            # Skip if we've already filled this position due to previous cell's span
            while col_idx < num_cols and grid[row_idx][col_idx]:
                col_idx += 1

            if col_idx >= num_cols:
                break

            # Get cell properties
            try:
                tc = cell._tc
                tc_pr = tc.tcPr

                # Grid span (horizontal merge)
                grid_span = tc_pr.gridSpan if tc_pr is not None else None
                h_span = int(grid_span.val) if grid_span is not None else 1

                # Vertical merge
                v_merge = tc_pr.vMerge if tc_pr is not None else None

            except Exception:
                h_span = 1
                v_merge = None

            # Get cell text and style
            cell_text = _get_cell_text(cell)
            cell_style = _extract_cell_style(cell)

            # Handle vertical merge
            if v_merge is not None:
                # Check if this is a continue (val is None or "continue")
                val = v_merge.val if hasattr(v_merge, "val") else None
                if val is None or val == "continue":
                    # This cell continues from the merge start above
                    # Use empty string (the real value is in the start cell)
                    for span_offset in range(h_span):
                        if col_idx + span_offset < num_cols:
                            grid[row_idx][col_idx + span_offset] = ""
                            styles[row_idx][col_idx + span_offset] = cell_style
                    col_idx += h_span
                    continue
                else:
                    # This is a merge restart - treat as new cell
                    vmerge_start[col_idx] = (row_idx, cell_text, cell_style)

            # Fill grid cells for this cell's span
            for span_offset in range(h_span):
                if col_idx + span_offset < num_cols:
                    if span_offset == 0:
                        grid[row_idx][col_idx] = cell_text
                        styles[row_idx][col_idx] = cell_style
                    else:
                        # Spanned cells get empty string
                        grid[row_idx][col_idx + span_offset] = ""
                        styles[row_idx][col_idx + span_offset] = cell_style

            col_idx += h_span

    return grid, styles


def _looks_numeric(value: str) -> bool:
    """Return True if *value* parses as a numeric data cell.

    Handles comma-decimal notation (``1234,5``) and NBSP/space thousand
    separators.  Does NOT exclude 4-digit years — callers that need
    year filtering (e.g. :func:`_is_header_like_row`) must do so
    themselves.
    """
    cleaned = value.replace(",", ".").replace("\u00a0", "").replace(" ", "")
    if not cleaned:
        return False
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def _classify_data_columns(
    data_rows: list[list[str]], num_cols: int
) -> list[str]:
    """Classify each column as ``"text"`` or ``"numeric"``.

    A column is ``"numeric"`` when >50 % of its non-empty cells parse as
    numbers (via :func:`_looks_numeric`).  Columns with no non-empty cells
    default to ``"numeric"`` (conservative — prevents false index detection).
    """
    result: list[str] = []
    for ci in range(num_cols):
        non_empty = 0
        numeric = 0
        for row in data_rows:
            val = (row[ci].strip() if ci < len(row) else "")
            if not val:
                continue
            non_empty += 1
            if _looks_numeric(val):
                numeric += 1
        if non_empty == 0:
            result.append("numeric")
        else:
            result.append("numeric" if numeric / non_empty > 0.5 else "text")
    return result


def _is_header_like_row(row: list[str]) -> bool:
    """Check if a row looks like a sub-header (mostly text, not data).

    Counts how many non-empty cells parse as numeric values. Cells that are
    exactly 4-digit years (e.g. "2025") are NOT counted as numeric because
    they commonly appear in sub-header rows as year labels.

    Returns True if fewer than 50% of non-empty cells are numeric.
    """
    non_empty = [c.strip() for c in row if c.strip()]
    if not non_empty:
        return False
    numeric_count = 0
    for cell in non_empty:
        if _looks_numeric(cell):
            # 4-digit integers are likely year labels, not data
            cleaned = cell.replace(",", ".").replace("\u00a0", "").replace(" ", "")
            if len(cleaned) == 4 and cleaned.isdigit():
                continue
            numeric_count += 1
    return numeric_count / len(non_empty) < 0.5


def _detect_header_rows_from_merges(table: "Table", *, max_scan: int = 10) -> int:
    """Detect how many leading rows are headers using merge information.

    A row is considered a header row if it contains cells with gridSpan > 1
    (horizontal spanning) or vMerge (vertical merge). These indicate
    structural headers with hierarchical labels.

    Scans the first *max_scan* rows and finds the last row with merge
    indicators. All rows from 0 up to (and including) that row are treated
    as header rows. This handles cases where a title row (no merges) precedes
    metric/year rows that do have merges.

    Returns at least 1 if the table has rows.
    """
    num_rows = len(table.rows)
    if num_rows == 0:
        return 0

    last_merge_row = -1
    for row_idx, row in enumerate(table.rows[: max_scan]):
        seen: set[int] = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen:
                continue
            seen.add(tc_id)
            try:
                tc = cell._tc
                tc_pr = tc.tcPr
                if tc_pr is not None:
                    gs = tc_pr.gridSpan
                    if gs is not None and int(gs.val) > 1:
                        last_merge_row = row_idx
                        break
                    vm = tc_pr.vMerge
                    if vm is not None:
                        last_merge_row = row_idx
                        break
            except Exception:
                continue

    # All rows up to and including the last merge row are headers
    header_count = last_merge_row + 1 if last_merge_row >= 0 else 0
    return max(header_count, 1) if num_rows > 0 else 0


def _build_compound_headers(
    header_rows: list[list[str]],
    *,
    title_row: str | None = None,
    data_rows: list[list[str]] | None = None,
) -> list[str]:
    """Stack header rows into compound column names joined by " / ".

    For each column position, join non-empty values from each header row.
    Empty cells in a row that are part of a horizontal span inherit the
    preceding non-empty cell's text (forward-fill within each row).

    When *data_rows* is provided, the forward-fill is **boundary-aware**:
    fill does not bleed from text index columns (leftmost contiguous text
    columns) into numeric data columns.  This prevents a merged "Region"
    cell from polluting data column headers.

    Args:
        header_rows: List of rows (each a list of cell strings).
        title_row: Optional title text prepended to all headers.
        data_rows: Optional data rows used to classify columns.

    Returns:
        List of compound header strings, one per column.
    """
    if not header_rows:
        return []

    num_cols = max(len(r) for r in header_rows) if header_rows else 0
    if num_cols == 0:
        return []

    # Identify index columns: leftmost contiguous text-data columns
    index_cols: set[int] = set()
    if data_rows is not None:
        col_types = _classify_data_columns(data_rows, num_cols)
        for ci in range(num_cols):
            if col_types[ci] == "text":
                index_cols.add(ci)
            else:
                break

    # Forward-fill each header row: empty cells inherit preceding non-empty
    filled_rows: list[list[str]] = []
    for row in header_rows:
        filled: list[str] = []
        last_nonempty = ""
        source_col = -1
        for i in range(num_cols):
            val = row[i].strip() if i < len(row) else ""
            if val:
                last_nonempty = val
                source_col = i
                filled.append(val)
            else:
                # Stop fill that crosses text index → numeric data boundary
                if source_col in index_cols and i not in index_cols:
                    filled.append("")
                else:
                    filled.append(last_nonempty)
        filled_rows.append(filled)

    # Stack columns vertically
    result: list[str] = []
    for col in range(num_cols):
        parts: list[str] = []
        if title_row:
            parts.append(title_row.strip())
        for row in filled_rows:
            val = row[col] if col < len(row) else ""
            if val and val not in parts:
                parts.append(val)
        result.append(" / ".join(parts) if parts else "")

    return result


# ---------------------------------------------------------------------------
# Visual analysis
# ---------------------------------------------------------------------------


def _analyze_visual_structure(
    styles: list[list[DocxCellStyle]],
) -> DocxVisualInfo:
    """Analyze visual structure of the table."""
    visual = DocxVisualInfo()

    if not styles or not styles[0]:
        return visual

    # Check for header fills in first few rows
    row_fills: list[tuple[float, float, float] | None] = []
    for ri, row in enumerate(styles[:5]):
        fills = [s.shading_color for s in row if s.shading_color is not None]
        if fills:
            row_fills.append(fills[0])
        else:
            row_fills.append(None)

    # Find header rows with consistent fill
    if row_fills and row_fills[0] is not None:
        first_fill = row_fills[0]
        for ri, fill in enumerate(row_fills):
            if fill == first_fill:
                visual.header_fill_rows.append(ri)
            else:
                break

    return visual


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract_tables_from_docx(
    docx_path: str | Path,
    *,
    extract_styles: bool = True,
) -> list[StructuredTable]:
    """Extract structured tables from a Word document.

    Args:
        docx_path: Path to the .docx file
        extract_styles: If True, extract visual style information.

    Returns:
        List of StructuredTable objects, one per table found.

    Raises:
        ImportError: If python-docx is not installed.
        FileNotFoundError: If the file doesn't exist.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install with: pip install pdf-ocr[docx]"
        ) from e

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    doc: Document = Document(path)
    result: list[StructuredTable] = []

    for table_idx, table in enumerate(doc.tables):
        # Extract grid
        grid, styles = _build_grid_from_table(table)

        if not grid or not grid[0]:
            continue

        # Normalize grid
        grid = normalize_grid(grid)

        # Analyze visual structure
        visual = _analyze_visual_structure(styles) if extract_styles else None

        # Estimate header rows
        header_count = 0
        if visual and visual.header_fill_rows:
            header_count = len(visual.header_fill_rows)
        else:
            header_count = estimate_header_rows(grid)

        # Split headers from data
        header_rows, data_rows = split_headers_from_data(grid, header_count)

        # Build column names
        column_names = build_column_names_from_headers(header_rows)

        # Ensure column names list has correct length
        num_cols = len(grid[0]) if grid else 0
        while len(column_names) < num_cols:
            column_names.append("")

        # Create metadata
        metadata = TableMetadata(
            page_number=0,  # DOCX doesn't have clear page boundaries
            table_index=table_idx,
            section_label=None,
        )

        # Create StructuredTable
        result.append(StructuredTable(
            column_names=column_names,
            data=data_rows,
            source_format="docx",
            metadata={
                "page_number": metadata.page_number,
                "table_index": metadata.table_index,
            },
        ))

    return result


# ---------------------------------------------------------------------------
# DOC (Legacy Word 97-2003) extraction via conversion
# ---------------------------------------------------------------------------


def _find_libreoffice() -> str | None:
    """Find LibreOffice executable on the system."""
    # Common paths for LibreOffice
    candidates = [
        "libreoffice",
        "soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows
    ]

    for candidate in candidates:
        if shutil.which(candidate):
            return candidate

    return None


def _convert_doc_to_docx(doc_path: Path, output_dir: Path) -> Path | None:
    """Convert a .doc file to .docx using LibreOffice.

    Returns the path to the converted file, or None if conversion fails.
    """
    libreoffice = _find_libreoffice()
    if not libreoffice:
        return None

    try:
        # Run LibreOffice in headless mode to convert
        result = subprocess.run(
            [
                libreoffice,
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(output_dir),
                str(doc_path),
            ],
            capture_output=True,
            timeout=60,  # 1 minute timeout
        )

        if result.returncode != 0:
            return None

        # Find the converted file
        expected_name = doc_path.stem + ".docx"
        converted_path = output_dir / expected_name

        if converted_path.exists():
            return converted_path

        return None

    except (subprocess.TimeoutExpired, subprocess.SubprocessError):
        return None


def _extract_tables_from_doc(
    doc_path: str | Path,
    *,
    extract_styles: bool = True,
) -> list[StructuredTable]:
    """Extract tables from a legacy .doc file by converting to .docx first.

    Requires LibreOffice to be installed for conversion.
    """
    path = Path(doc_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {doc_path}")

    # Check if LibreOffice is available
    if not _find_libreoffice():
        raise ImportError(
            "LibreOffice is required to read legacy .doc files. "
            "Please install LibreOffice or convert the file to .docx format. "
            "Download: https://www.libreoffice.org/download/"
        )

    # Convert in a temporary directory
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        converted = _convert_doc_to_docx(path, tmpdir_path)

        if not converted:
            raise RuntimeError(
                f"Failed to convert {path.name} to .docx format. "
                "Please convert the file manually or check your LibreOffice installation."
            )

        # Extract from the converted file
        result = extract_tables_from_docx(converted, extract_styles=extract_styles)

        # Update source format in metadata
        for table in result:
            table.source_format = "doc"

        return result


def extract_tables_from_word(
    word_path: str | Path,
    *,
    extract_styles: bool = True,
) -> list[StructuredTable]:
    """Extract structured tables from a Word document (DOC or DOCX).

    Auto-detects format based on file extension and uses the appropriate
    extraction method.

    Args:
        word_path: Path to the Word file (.docx or .doc)
        extract_styles: If True, extract visual style information.

    Returns:
        List of StructuredTable objects.

    Note: .doc files require LibreOffice for conversion.
    """
    path = Path(word_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return extract_tables_from_docx(path, extract_styles=extract_styles)
    elif suffix == ".doc":
        return _extract_tables_from_doc(path, extract_styles=extract_styles)
    else:
        # Try DOCX first (more common), fall back to DOC
        try:
            return extract_tables_from_docx(path, extract_styles=extract_styles)
        except Exception:
            return _extract_tables_from_doc(path, extract_styles=extract_styles)


def docx_to_markdown(docx_path: str | Path) -> str:
    """Extract tables from Word and render as markdown.

    Convenience function for quick inspection of Word content.
    """
    tables = extract_tables_from_docx(docx_path, extract_styles=False)

    parts: list[str] = []
    for idx, table in enumerate(tables):
        parts.append(f"## Table {idx + 1}\n")

        if table.column_names:
            parts.append("| " + " | ".join(table.column_names) + " |")
            parts.append("| " + " | ".join(["---"] * len(table.column_names)) + " |")

        for row in table.data:
            parts.append("| " + " | ".join(row) + " |")

        parts.append("")  # Blank line between tables

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# DOCX → Compressed markdown pipeline
# ---------------------------------------------------------------------------


def compress_docx_table(
    raw_grid: list[list[str]],
    header_row_count: int,
    *,
    title: str | None = None,
) -> str:
    """Render a DOCX table as pipe-table markdown with compound headers.

    Splits raw_grid into header rows (0..header_row_count-1) and data rows,
    builds compound "/" headers with forward-fill for spans, strips trailing
    empty columns, and renders as markdown pipe-table.

    Args:
        raw_grid: The full grid including header and data rows.
        header_row_count: Number of leading rows to treat as headers.
        title: Optional section title prepended as ``## title``.

    Returns:
        Markdown string with optional title and pipe-table.
    """
    if not raw_grid:
        return ""

    header_rows = raw_grid[:header_row_count]
    data_rows = raw_grid[header_row_count:]

    # Build compound headers
    headers = _build_compound_headers(header_rows, data_rows=data_rows)
    if not headers:
        return ""

    num_cols = len(headers)

    # Strip trailing empty columns
    while num_cols > 0:
        col = num_cols - 1
        header_empty = not headers[col]
        data_empty = all(
            not (row[col] if col < len(row) else "")
            for row in data_rows
        )
        if header_empty and data_empty:
            num_cols -= 1
        else:
            break
    if num_cols == 0:
        return ""

    headers = headers[:num_cols]

    parts: list[str] = []

    if title:
        parts.append(f"## {title.strip()}")
        parts.append("")

    # Header row
    parts.append("| " + " | ".join(headers) + " |")
    # Separator row
    parts.append("| " + " | ".join(["---"] * num_cols) + " |")
    # Data rows
    for row in data_rows:
        cells = [(row[i] if i < len(row) else "") for i in range(num_cols)]
        parts.append("| " + " | ".join(cells) + " |")

    return "\n".join(parts)


def compress_docx_tables(
    docx_path: str | Path,
    *,
    table_indices: list[int] | None = None,
) -> list[tuple[str, dict]]:
    """Extract tables from a DOCX and render each as compressed pipe-table markdown.

    End-to-end pipeline: extract DOCX → build raw grid (with _tc dedup) →
    detect header rows via merge info → build compound headers → render
    pipe-table markdown.

    Args:
        docx_path: Path to the .docx file.
        table_indices: If provided, only process tables at these indices.

    Returns:
        List of ``(markdown_text, metadata)`` tuples. Metadata includes
        ``table_index``, ``title``, ``row_count``, ``col_count``.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install with: pip install pdf-ocr[docx]"
        ) from e

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    doc = Document(path)
    results: list[tuple[str, dict]] = []

    for table_idx, table in enumerate(doc.tables):
        if table_indices is not None and table_idx not in table_indices:
            continue

        grid, _styles = _build_grid_from_table(table)
        if not grid or not grid[0]:
            continue

        # Detect header rows from merge information
        header_count = _detect_header_rows_from_merges(table)

        # Content-based extension: if the row right after merge-detected
        # headers looks like a sub-header (mostly text/labels) but the row
        # after that is clearly data (mostly numeric), bump header_count.
        # This catches tables where row 0 has group spans ("spring crops")
        # and row 1 has metric labels ("MOA Target 2025") without merges.
        if (
            header_count >= 1
            and header_count < len(grid)
            and header_count + 1 < len(grid)
            and _is_header_like_row(grid[header_count])
            and not _is_header_like_row(grid[header_count + 1])
        ):
            header_count += 1

        # Detect title: if row 0 has only one non-empty cell, treat as title
        title: str | None = None
        non_empty = [c for c in grid[0] if c.strip()]
        if len(non_empty) == 1 and header_count > 1:
            title = non_empty[0]
            grid = grid[1:]  # remove title row from grid
            header_count -= 1

        data_row_count = len(grid) - header_count

        md = compress_docx_table(grid, header_count, title=title)
        if not md:
            continue

        # Determine column count from first data row (or headers)
        col_count = len(grid[0]) if grid else 0

        meta = {
            "table_index": table_idx,
            "title": title,
            "row_count": data_row_count,
            "col_count": col_count,
        }
        results.append((md, meta))

    return results


# ---------------------------------------------------------------------------
# Table classification
# ---------------------------------------------------------------------------


def classify_docx_tables(
    docx_path: str | Path,
    categories: dict[str, list[str]],
    *,
    min_data_rows: int = 0,
    propagate: bool = False,
    propagate_threshold: float = 0.3,
) -> list[dict]:
    """Classify each table in a DOCX by matching header text against user-defined keywords.

    Thin wrapper: compresses tables via ``compress_docx_tables()`` then
    delegates to the format-agnostic ``classify_tables()``.

    Args:
        docx_path: Path to the .docx file.
        categories: Mapping of category name → list of keywords.
            Each table is scored against every category by counting how many
            keywords appear (using word-boundary matching) in the combined
            title + header text (case-insensitive). The table is assigned to
            the highest-scoring category. Ties are broken by dict iteration
            order. Tables that match no category are labelled ``"other"``.

            Example::

                categories = {
                    "harvest": ["area harvested", "yield", "collected"],
                    "export": ["export", "shipment", "fob"],
                }

        min_data_rows: Tables with fewer data rows are forced to ``"other"``
            and excluded from similarity profile building.
        propagate: If True, after keyword scoring, build structural profiles
            from matched tables and propagate categories to unmatched ones
            that exceed *propagate_threshold* similarity.
        propagate_threshold: Minimum similarity score (0-1) for propagation.

    Returns:
        List of dicts with ``index``, ``category``, ``title``, ``rows``, ``cols``.
    """
    from pdf_ocr.classify import classify_tables

    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Install with: pip install pdf-ocr[docx]"
        ) from e

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    # Count total tables so we can include empty ones as "other"
    doc = Document(path)
    total_tables = len(doc.tables)

    compressed = compress_docx_tables(docx_path)
    results = classify_tables(
        compressed,
        categories,
        min_data_rows=min_data_rows,
        propagate=propagate,
        propagate_threshold=propagate_threshold,
    )

    # compress_docx_tables may skip tables with empty grids.
    # Re-insert them as "other" to match the old per-table behaviour.
    if len(results) < total_tables:
        present = {r["index"] for r in results}
        extras = []
        for idx in range(total_tables):
            if idx not in present:
                extras.append({
                    "index": idx,
                    "category": "other",
                    "title": None,
                    "rows": 0,
                    "cols": 0,
                })
        results = sorted(results + extras, key=lambda r: r["index"])

    return results


# ---------------------------------------------------------------------------
# Pivot value extraction
# ---------------------------------------------------------------------------

from pdf_ocr.classify import _YEAR_RE


def extract_pivot_values(markdown: str) -> list[str]:
    """Extract pivot column values (years, periods) from compressed markdown headers.

    Scans all pipe-table header lines (before the ``|---|`` separator) for
    4-digit years (1900-2099). Returns them sorted in ascending order,
    deduplicated.

    Use this to dynamically build schema aliases for the ``year`` column
    instead of hardcoding years::

        results = compress_docx_tables("report.docx", table_indices=[4])
        md, meta = results[0]
        years = extract_pivot_values(md)          # e.g. ["2024", "2025"]
        recent = years[-2:]                       # last 2 years
        schema_dict["columns"][2]["aliases"] = recent

    Args:
        markdown: Compressed pipe-table markdown from ``compress_docx_table()``.

    Returns:
        Sorted list of unique year strings found in the header row.
    """
    years: set[str] = set()
    for line in markdown.split("\n"):
        if not line.startswith("|"):
            continue
        if "---" in line:
            break  # separator = end of headers
        years.update(_YEAR_RE.findall(line))
    return sorted(years)
