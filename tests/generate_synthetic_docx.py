#!/usr/bin/env python3
"""Generate synthetic DOCX test files in inputs/docx/synthetic/.

These files are used by tests/test_docx_extractor.py (section 12).
Run:  uv run python tests/generate_synthetic_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

OUT = Path(__file__).resolve().parent.parent / "inputs" / "docx" / "synthetic"


def build(name: str, table_specs: list[dict]) -> None:
    doc = Document()
    for spec in table_specs:
        grid = spec["grid"]
        nrows = len(grid)
        ncols = max(len(r) for r in grid) if grid else 0
        tbl = doc.add_table(rows=nrows, cols=ncols)
        tbl.style = "Table Grid"

        skip: set[tuple[int, int]] = set()
        for row, c1, c2 in spec.get("hmerges", []):
            tbl.cell(row, c1).merge(tbl.cell(row, c2))
            for c in range(c1 + 1, c2 + 1):
                skip.add((row, c))
        for col, r1, r2 in spec.get("vmerges", []):
            tbl.cell(r1, col).merge(tbl.cell(r2, col))
            for r in range(r1 + 1, r2 + 1):
                skip.add((r, col))

        for ri, row_data in enumerate(grid):
            for ci, val in enumerate(row_data):
                if ci < ncols and (ri, ci) not in skip and val:
                    tbl.cell(ri, ci).text = val

        doc.add_paragraph()

    OUT.mkdir(parents=True, exist_ok=True)
    p = OUT / f"{name}.docx"
    doc.save(p)
    print(f"  {p.name}")


# -----------------------------------------------------------------
# 1. flat.docx — no merges, simple 3-column table
# -----------------------------------------------------------------
build("flat", [{
    "grid": [
        ["Name", "Age", "City"],
        ["Alice", "30", "Paris"],
        ["Bob", "25", "London"],
        ["Carol", "35", "Berlin"],
        ["Diana", "28", "Tokyo"],
        ["Eve", "42", "Sydney"],
    ],
}])

# -----------------------------------------------------------------
# 2. hspan.docx — horizontal spans producing compound headers
# -----------------------------------------------------------------
build("hspan", [{
    "grid": [
        ["Category", "Revenue", "", "Cost", ""],
        ["", "Q1", "Q2", "Q1", "Q2"],
        ["Product A", "100", "120", "80", "90"],
        ["Product B", "200", "220", "150", "160"],
        ["Product C", "350", "400", "290", "310"],
        ["Product D", "500", "480", "410", "390"],
    ],
    "hmerges": [(0, 1, 2), (0, 3, 4)],
    "vmerges": [(0, 0, 1)],
}])

# -----------------------------------------------------------------
# 3. title_hspan.docx — title row + hierarchical headers
# -----------------------------------------------------------------
build("title_hspan", [{
    "grid": [
        ["QUARTERLY REPORT", "", "", "", ""],
        ["Region", "Revenue", "", "Cost", ""],
        ["", "Q1", "Q2", "Q1", "Q2"],
        ["North", "1000", "1100", "800", "850"],
        ["South", "900", "950", "700", "720"],
        ["East", "1200", "1300", "1000", "1050"],
        ["West", "800", "780", "650", "640"],
    ],
    "hmerges": [(0, 0, 4), (1, 1, 2), (1, 3, 4)],
    "vmerges": [(0, 1, 2)],
}])

# -----------------------------------------------------------------
# 4. deep_hierarchy.docx — 3-level header hierarchy
# -----------------------------------------------------------------
build("deep_hierarchy", [{
    "grid": [
        ["ID", "Group A", "", "", "Group B", "", ""],
        ["", "Sub X", "", "Sub Y", "Sub X", "", "Sub Y"],
        ["", "2024", "2025", "2024", "2024", "2025", "2024"],
        ["Region 1", "10", "15", "8", "12", "18", "9"],
        ["Region 2", "20", "25", "18", "22", "28", "19"],
        ["Region 3", "30", "35", "28", "32", "38", "29"],
        ["Region 4", "40", "45", "38", "42", "48", "39"],
        ["Region 5", "50", "55", "48", "52", "58", "49"],
    ],
    "hmerges": [(0, 1, 3), (0, 4, 6), (1, 1, 2), (1, 4, 5)],
    "vmerges": [(0, 0, 2)],
}])

# -----------------------------------------------------------------
# 5. multi_category.docx — multiple tables for classification
# -----------------------------------------------------------------
build("multi_category", [
    # Table 0: shipping/logistics
    {
        "grid": [
            ["Port", "Cargo volume", "Vessel count"],
            ["Rotterdam", "50000", "120"],
            ["Shanghai", "80000", "200"],
            ["Singapore", "65000", "180"],
        ],
    },
    # Table 1: HR / payroll
    {
        "grid": [
            ["Employee", "Department", "Salary", "Start date"],
            ["Alice", "Engineering", "120k", "2022-01-15"],
            ["Bob", "Marketing", "95k", "2021-06-01"],
            ["Carol", "Finance", "110k", "2020-03-20"],
        ],
    },
    # Table 2: financial summary with title + merges
    {
        "grid": [
            ["FINANCIAL SUMMARY", "", "", "", ""],
            ["Metric", "Revenue", "", "Expenses", ""],
            ["", "2024", "2025", "2024", "2025"],
            ["Q1", "1500", "1700", "1200", "1350"],
            ["Q2", "1600", "1800", "1250", "1400"],
            ["Q3", "1700", "1900", "1300", "1450"],
            ["Q4", "1800", "2100", "1350", "1500"],
        ],
        "hmerges": [(0, 0, 4), (1, 1, 2), (1, 3, 4)],
        "vmerges": [(0, 1, 2)],
    },
    # Table 3: inventory/stock
    {
        "grid": [
            ["Item", "Stock level", "Warehouse", "Reorder point"],
            ["Widget A", "5000", "NYC", "1000"],
            ["Widget B", "3000", "LA", "800"],
            ["Widget C", "7500", "Chicago", "1500"],
        ],
    },
])

# -----------------------------------------------------------------
# 6. wide_table.docx — 10 columns with 2 span groups
# -----------------------------------------------------------------
build("wide_table", [{
    "grid": [
        ["ID", "Block A", "", "", "", "Block B", "", "", "", ""],
        ["", "C1", "C2", "C3", "C4", "C5", "C6", "C7", "C8", "C9"],
        ["R1", "a", "b", "c", "d", "e", "f", "g", "h", "i"],
        ["R2", "j", "k", "l", "m", "n", "o", "p", "q", "r"],
        ["R3", "s", "t", "u", "v", "w", "x", "y", "z", "A"],
    ],
    "hmerges": [(0, 1, 4), (0, 5, 9)],
    "vmerges": [(0, 0, 1)],
}])

# -----------------------------------------------------------------
# 7. unicode.docx — non-ASCII content
# -----------------------------------------------------------------
build("unicode", [{
    "grid": [
        ["Région", "Données", "", ""],
        ["", "2023", "2024", "2025"],
        ["Île-de-France", "1 200", "1 500", "1 600"],
        ["Rhône-Alpes", "1 800", "2 000", "2 100"],
        ["Provence-Alpes-Côte d'Azur", "950", "1 050", "1 100"],
        ["Occitanie", "1 100", "1 250", "1 350"],
    ],
    "hmerges": [(0, 1, 3)],
    "vmerges": [(0, 0, 1)],
}])

# -----------------------------------------------------------------
# 8. single_column.docx — minimal edge case
# -----------------------------------------------------------------
build("single_column", [{
    "grid": [
        ["Countries"],
        ["France"],
        ["Germany"],
        ["Italy"],
        ["Spain"],
    ],
}])

# -----------------------------------------------------------------
# 9. propagation.docx — for similarity propagation tests
#    Table 0: keyword-matchable (8 cols, "area harvested", "yield")
#    Table 1: same structure but no keywords (8 cols, overlapping tokens)
#    Table 2: different structure (3 cols, unrelated tokens)
# -----------------------------------------------------------------
build("propagation", [
    # Table 0: matches "area harvested" and "yield"
    {
        "grid": [
            ["Region", "Area harvested", "", "Yield", "", "Production", "", ""],
            ["", "2024", "2025", "2024", "2025", "2024", "2025", ""],
            ["North", "100", "110", "3.5", "3.8", "350", "418", ""],
            ["South", "200", "210", "4.0", "4.2", "800", "882", ""],
            ["East", "150", "160", "3.2", "3.5", "480", "560", ""],
        ],
        "hmerges": [(0, 1, 2), (0, 3, 4), (0, 5, 6)],
        "vmerges": [(0, 0, 1)],
    },
    # Table 1: same 8-col structure, overlapping tokens, but no keywords
    {
        "grid": [
            ["Region", "Sown acreage", "", "Output rate", "", "Total volume", "", ""],
            ["", "2024", "2025", "2024", "2025", "2024", "2025", ""],
            ["North", "100", "110", "3.5", "3.8", "350", "418", ""],
            ["South", "200", "210", "4.0", "4.2", "800", "882", ""],
            ["East", "150", "160", "3.2", "3.5", "480", "560", ""],
        ],
        "hmerges": [(0, 1, 2), (0, 3, 4), (0, 5, 6)],
        "vmerges": [(0, 0, 1)],
    },
    # Table 2: different structure (3 cols, unrelated)
    {
        "grid": [
            ["Date", "Price", "Volume"],
            ["2025-01-01", "250", "1000"],
            ["2025-01-02", "255", "1200"],
        ],
    },
])

print(f"Done: {len(list(OUT.glob('*.docx')))} files in {OUT}")
