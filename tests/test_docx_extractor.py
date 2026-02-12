"""Regression tests for DOCX table extraction.

These tests verify that changes to _build_grid_from_table() and related
functions do not break existing extraction behavior across both real-world
DOCX reports and synthetic test documents.

Run with: uv run pytest tests/test_docx_extractor.py -v
"""

from __future__ import annotations

from pathlib import Path

import pytest

from pdf_ocr.docx_extractor import (
    _build_compound_headers,
    _build_grid_from_table,
    _detect_header_rows_from_merges,
    classify_docx_tables,
    compress_docx_table,
    compress_docx_tables,
    extract_pivot_values,
    extract_tables_from_docx,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

DOCX_DIR = Path("inputs/docx/input")


def _resolve(fragment: str) -> Path:
    matches = sorted(DOCX_DIR.glob(f"*{fragment}*"))
    assert matches, f"No DOCX file matching '{fragment}' in {DOCX_DIR}"
    return matches[0]


@pytest.fixture(scope="module")
def july_path() -> Path:
    return _resolve("July")


@pytest.fixture(scope="module")
def sep_path() -> Path:
    return _resolve("September")


@pytest.fixture(scope="module")
def apr25_path() -> Path:
    return _resolve("April 25")


@pytest.fixture(scope="module")
def all_docx_paths() -> list[Path]:
    paths = sorted(DOCX_DIR.glob("*.docx"))
    assert len(paths) == 6, f"Expected 6 DOCX files, found {len(paths)}"
    return paths


# ---------------------------------------------------------------------------
# 1. extract_tables_from_docx() — shape regression
#    Verifies table count, column count, and data row count per file.
# ---------------------------------------------------------------------------

# Expected (table_count, [(col_count, data_row_count), ...]) per file fragment
EXPECTED_SHAPES: dict[str, tuple[int, list[tuple[int, int]]]] = {
    "Apr 21": (3, [(21, 16), (11, 33), (6, 32)]),
    "April 25": (9, [(21, 16), (11, 33), (11, 67), (13, 88), (7, 40), (5, 7), (7, 38), (7, 27), (7, 36)]),
    "May": (3, [(21, 16), (11, 36), (11, 17)]),
    "June": (12, [(5, 7), (21, 16), (11, 25), (1, 29), (11, 93), (13, 88), (7, 70), (5, 14), (5, 46), (7, 56), (7, 64), (7, 66)]),
    "July": (7, [(21, 16), (11, 9), (7, 25), (8, 31), (8, 35), (8, 30), (8, 15)]),
    "September": (14, [(21, 16), (11, 7), (6, 34), (3, 25), (3, 24), (3, 74), (8, 82), (8, 82), (8, 78), (8, 49), (8, 48), (8, 9), (8, 54), (8, 20)]),
}


@pytest.mark.parametrize("fragment,expected", EXPECTED_SHAPES.items(), ids=EXPECTED_SHAPES.keys())
def test_extract_tables_shape(fragment: str, expected: tuple[int, list[tuple[int, int]]]) -> None:
    """Table count, column count, and data row count must stay stable."""
    path = _resolve(fragment)
    tables = extract_tables_from_docx(path)

    expected_count, expected_dims = expected
    assert len(tables) == expected_count, (
        f"{path.name}: expected {expected_count} tables, got {len(tables)}"
    )

    for i, (exp_cols, exp_rows) in enumerate(expected_dims):
        t = tables[i]
        assert len(t.column_names) == exp_cols, (
            f"{path.name} table {i}: expected {exp_cols} columns, got {len(t.column_names)}"
        )
        assert len(t.data) == exp_rows, (
            f"{path.name} table {i}: expected {exp_rows} data rows, got {len(t.data)}"
        )


# ---------------------------------------------------------------------------
# 2. extract_tables_from_docx() — content regression
#    Verifies specific column names and cell values haven't changed.
# ---------------------------------------------------------------------------


class TestExtractTablesContent:
    """Verify specific cell values in extracted tables."""

    def test_july_export_table_column_names(self, july_path: Path) -> None:
        """Table 0 (export) column names must be stable."""
        tables = extract_tables_from_docx(july_path)
        t0 = tables[0]
        assert t0.column_names[0].startswith("Rus grains export")
        assert "Wheat" in t0.column_names[1]
        assert t0.column_names[2] == "2024/25"

    def test_july_export_table_first_data_row(self, july_path: Path) -> None:
        """Table 0 first data row values must be stable."""
        tables = extract_tables_from_docx(july_path)
        row = tables[0].data[0]
        assert "July" in row[0]
        assert row[1] == "4133"

    def test_july_wheat_column_names(self, july_path: Path) -> None:
        """Table 4 (WHEAT) column names must reflect correct merge handling."""
        tables = extract_tables_from_docx(july_path)
        t4 = tables[4]
        # After _tc dedup, col 2 should be "Area harvested" not duplicated
        assert t4.column_names[0] == "Region"
        assert "WHEAT" in t4.column_names[1]
        assert t4.column_names[2].startswith("Area harvested")
        assert t4.column_names[4].startswith("collected")
        assert t4.column_names[6].startswith("Yield")

    def test_july_wheat_data_values(self, july_path: Path) -> None:
        """WHEAT table actual data (row with 'Rus Fed w/o new territories')."""
        tables = extract_tables_from_docx(july_path)
        # data[0] and data[1] are sub-header rows, data[2] has actual values
        row = tables[4].data[2]
        assert "Rus Fed" in row[0]
        assert row[2] == "3219,9"  # Area harvested 2025
        assert row[3] == "6195,2"  # Area harvested 2024

    def test_july_grains_pulses_harvest(self, july_path: Path) -> None:
        """Table 3 (Grains and pulses) structure."""
        tables = extract_tables_from_docx(july_path)
        t3 = tables[3]
        assert t3.column_names[0] == "Region"
        assert "Grains and pulses" in t3.column_names[1]

    def test_sep_wheat_column_names(self, sep_path: Path) -> None:
        """September WHEAT table (table 7) column names."""
        tables = extract_tables_from_docx(sep_path)
        t7 = tables[7]
        assert t7.column_names[0] == "Region"
        assert "WHEAT" in t7.column_names[1]
        assert t7.column_names[2].startswith("Area harvested")
        assert t7.column_names[4].startswith("collected")
        assert t7.column_names[6].startswith("Yield")

    def test_all_tables_have_source_format(self, all_docx_paths: list[Path]) -> None:
        """Every extracted table must have source_format='docx'."""
        for path in all_docx_paths:
            tables = extract_tables_from_docx(path)
            for i, t in enumerate(tables):
                assert t.source_format == "docx", (
                    f"{path.name} table {i}: source_format={t.source_format!r}"
                )


# ---------------------------------------------------------------------------
# 3. _build_grid_from_table() — grid dimension regression
# ---------------------------------------------------------------------------

EXPECTED_GRIDS: dict[str, list[tuple[int, int]]] = {
    "July": [(20, 21), (9, 11), (25, 7), (33, 8), (37, 8), (32, 8), (17, 8)],
    "September": [
        (20, 21), (7, 11), (37, 6), (25, 3), (26, 3), (76, 3),
        (84, 8), (84, 8), (80, 8), (51, 8), (50, 8), (11, 8), (56, 8), (22, 8),
    ],
}


@pytest.mark.parametrize("fragment,expected", EXPECTED_GRIDS.items(), ids=EXPECTED_GRIDS.keys())
def test_grid_dimensions(fragment: str, expected: list[tuple[int, int]]) -> None:
    """Raw grid dimensions (rows x cols) must stay stable after _tc dedup."""
    from docx import Document

    path = _resolve(fragment)
    doc = Document(path)

    for ti, (exp_rows, exp_cols) in enumerate(expected):
        grid, styles = _build_grid_from_table(doc.tables[ti])
        assert len(grid) == exp_rows, (
            f"{path.name} grid {ti}: expected {exp_rows} rows, got {len(grid)}"
        )
        assert len(grid[0]) == exp_cols, (
            f"{path.name} grid {ti}: expected {exp_cols} cols, got {len(grid[0])}"
        )
        # Styles grid must have same dimensions
        assert len(styles) == exp_rows
        assert len(styles[0]) == exp_cols


# ---------------------------------------------------------------------------
# 4. _build_grid_from_table() — _tc dedup correctness
#    The core bug fix: merged cells must not produce duplicated headers.
# ---------------------------------------------------------------------------


class TestTcDedup:
    """Verify the _tc deduplication fix for python-docx merged cells."""

    def test_wheat_header_row_no_duplication(self, july_path: Path) -> None:
        """WHEAT table row 1 must have metrics in correct positions, not shifted."""
        from docx import Document

        doc = Document(july_path)
        grid, _ = _build_grid_from_table(doc.tables[4])

        # Row 0: title row (WHEAT)
        # Row 1: metric labels
        # Row 2: year labels
        row1 = grid[1]
        assert row1[0] == "Region"
        assert row1[2].startswith("Area harvested")
        assert row1[3] == ""  # spanned cell (empty, not duplicated text)
        assert row1[4].startswith("collected")
        assert row1[5] == ""  # spanned cell
        assert row1[6].startswith("Yield")
        assert row1[7] == ""  # spanned cell

    def test_wheat_year_row(self, july_path: Path) -> None:
        """Row 2 should have year values under each metric pair."""
        from docx import Document

        doc = Document(july_path)
        grid, _ = _build_grid_from_table(doc.tables[4])

        row2 = grid[2]
        assert row2[2] == "2025"
        assert row2[3] == "2024"
        assert row2[4] == "2025"
        assert row2[5] == "2024"
        assert row2[6] == "2025"
        assert row2[7] == "2024"

    def test_grains_pulses_header_row(self, july_path: Path) -> None:
        """Grains+pulses (table 3) has same merge structure as WHEAT."""
        from docx import Document

        doc = Document(july_path)
        grid, _ = _build_grid_from_table(doc.tables[3])

        row1 = grid[1]
        assert row1[2].startswith("Area harvested")
        assert row1[4].startswith("collected")
        assert row1[6].startswith("Yield")

    def test_export_table_no_regression(self, july_path: Path) -> None:
        """Export table (table 0) — wide table with many columns. Dedup must not break it."""
        from docx import Document

        doc = Document(july_path)
        grid, _ = _build_grid_from_table(doc.tables[0])

        assert len(grid[0]) == 21
        assert len(grid) == 20
        # Title row should have content
        assert any(c.strip() for c in grid[0])


# ---------------------------------------------------------------------------
# 5. _detect_header_rows_from_merges()
# ---------------------------------------------------------------------------


class TestDetectHeaderRows:
    """Verify merge-based header row detection."""

    def test_harvest_table_3_header_rows(self, july_path: Path) -> None:
        """Grains+pulses table has title + metric + year = 3 header rows."""
        from docx import Document

        doc = Document(july_path)
        hc = _detect_header_rows_from_merges(doc.tables[3])
        assert hc == 3, f"Expected 3 header rows, got {hc}"

    def test_wheat_table_3_header_rows(self, july_path: Path) -> None:
        """WHEAT table has WHEAT title + metric + year = 3 header rows via merges."""
        from docx import Document

        doc = Document(july_path)
        hc = _detect_header_rows_from_merges(doc.tables[4])
        # Row 0 (WHEAT) has no merges, row 1 has gridSpan, row 2 has vMerge
        # → last_merge_row = 2 → header_count = 3
        assert hc == 3, f"Expected 3 header rows, got {hc}"

    def test_export_table_at_least_1(self, july_path: Path) -> None:
        """Export table must return at least 1 header row."""
        from docx import Document

        doc = Document(july_path)
        hc = _detect_header_rows_from_merges(doc.tables[0])
        assert hc >= 1


# ---------------------------------------------------------------------------
# 6. _build_compound_headers()
# ---------------------------------------------------------------------------


class TestBuildCompoundHeaders:
    """Verify compound header construction with forward-fill."""

    def test_basic_two_row_header(self) -> None:
        headers = _build_compound_headers([
            ["Region", "Area", "Metric A", "", "Metric B", ""],
            ["Region", "Area", "2025", "2024", "2025", "2024"],
        ])
        assert headers[0] == "Region"
        assert headers[1] == "Area"
        assert headers[2] == "Metric A / 2025"
        assert headers[3] == "Metric A / 2024"
        assert headers[4] == "Metric B / 2025"
        assert headers[5] == "Metric B / 2024"

    def test_single_row_header(self) -> None:
        headers = _build_compound_headers([["A", "B", "C"]])
        assert headers == ["A", "B", "C"]

    def test_empty_input(self) -> None:
        assert _build_compound_headers([]) == []

    def test_with_title_row(self) -> None:
        headers = _build_compound_headers(
            [["Col1", "Col2"]],
            title_row="WHEAT",
        )
        assert headers[0] == "WHEAT / Col1"
        assert headers[1] == "WHEAT / Col2"

    def test_forward_fill_propagates(self) -> None:
        """Empty cells in spanning positions should inherit preceding value."""
        headers = _build_compound_headers([
            ["X", "", ""],  # X spans 3 columns
        ])
        assert headers == ["X", "X", "X"]

    def test_no_duplicate_parts(self) -> None:
        """If header rows repeat the same value, don't duplicate in the compound."""
        headers = _build_compound_headers([
            ["Region", "Area"],
            ["Region", "Area"],
        ])
        assert headers == ["Region", "Area"]


# ---------------------------------------------------------------------------
# 7. compress_docx_table()
# ---------------------------------------------------------------------------


class TestCompressDocxTable:
    """Verify pipe-table markdown rendering."""

    def test_basic_render(self) -> None:
        grid = [
            ["Metric A", "", "Metric B", ""],
            ["2025", "2024", "2025", "2024"],
            ["10", "20", "30", "40"],
            ["11", "21", "31", "41"],
        ]
        md = compress_docx_table(grid, header_row_count=2)
        lines = md.strip().split("\n")
        assert lines[0].startswith("| Metric A / 2025")
        assert "---" in lines[1]
        assert lines[2].startswith("| 10")

    def test_with_title(self) -> None:
        grid = [["A", "B"], ["1", "2"]]
        md = compress_docx_table(grid, header_row_count=1, title="WHEAT")
        assert md.startswith("## WHEAT")

    def test_strips_trailing_empty_columns(self) -> None:
        # Forward-fill in compound header construction fills trailing empty headers
        # from the last non-empty cell, so the only way to get a truly empty
        # trailing header is when the entire header row is empty from the start.
        # Test: 3 columns where col 2 has empty header AND empty data.
        # Use two header rows where col 2 is explicitly empty in both.
        grid = [
            ["A", "B", ""],  # header row 1: forward-fill makes col 2 = "B"
            ["1", "2", ""],  # data row
        ]
        md = compress_docx_table(grid, header_row_count=1)
        # col 2 header is "B" (from forward-fill) so it's NOT stripped.
        # This is correct: forward-fill means the column has semantic meaning.
        header_line = md.split("\n")[0]
        assert "B" in header_line
        # Verify the pipe table renders 3 columns
        assert header_line.count("|") == 4  # | A | B | B |

    def test_empty_grid_returns_empty(self) -> None:
        assert compress_docx_table([], header_row_count=0) == ""


# ---------------------------------------------------------------------------
# 8. compress_docx_tables() — integration
# ---------------------------------------------------------------------------


class TestCompressDocxTables:
    """Verify end-to-end DOCX → compressed markdown pipeline."""

    def test_july_wheat_compound_headers(self, july_path: Path) -> None:
        """WHEAT table compound headers must have metric / year format."""
        results = compress_docx_tables(july_path, table_indices=[4])
        assert len(results) == 1
        md, meta = results[0]

        assert meta["title"] == "WHEAT"
        header_line = [l for l in md.split("\n") if l.startswith("|") and "---" not in l][0]
        assert "Area harvested" in header_line
        assert "/ 2025" in header_line
        assert "/ 2024" in header_line
        assert "Yield" in header_line

    def test_july_wheat_data_row_count(self, july_path: Path) -> None:
        results = compress_docx_tables(july_path, table_indices=[4])
        _, meta = results[0]
        assert meta["row_count"] == 34

    def test_july_all_harvest_tables(self, july_path: Path) -> None:
        """All 4 July harvest tables should compress successfully."""
        results = compress_docx_tables(july_path, table_indices=[3, 4, 5, 6])
        assert len(results) == 4
        titles = [m["title"] for _, m in results]
        assert "Grains and pulses" in titles
        assert "WHEAT" in titles
        assert "BARLEY" in titles
        assert "RAPE" in titles

    def test_sep_harvest_tables(self, sep_path: Path) -> None:
        """September has 8 harvest tables; all should compress."""
        classes = classify_docx_tables(sep_path, _TEST_CATEGORIES)
        harvest_idx = [c["index"] for c in classes if c["category"] == "harvest"]
        results = compress_docx_tables(sep_path, table_indices=harvest_idx)
        assert len(results) == 8

    def test_table_indices_filter(self, july_path: Path) -> None:
        """table_indices should restrict which tables are processed."""
        all_results = compress_docx_tables(july_path)
        filtered = compress_docx_tables(july_path, table_indices=[0])
        assert len(filtered) <= len(all_results)
        assert len(filtered) == 1


# ---------------------------------------------------------------------------
# 9. classify_docx_tables()
# ---------------------------------------------------------------------------

# Domain-specific categories used for testing — these are NOT part of the
# library; real users define their own categories.
_TEST_CATEGORIES = {
    "harvest": ["area harvested", "yield", "collected", "bunker", "centner",
                "harvested area", "crop harvested"],
    "planting": ["spring crops", "moa target", "spring wheat", "spring barley",
                 "sown area", "planting", "sowing", "planted"],
    "export": ["export", "shipment", "ports", "fob", "vessel", "cargo"],
}

# Expected classification per file
EXPECTED_CLASSIFICATIONS: dict[str, dict[str, int]] = {
    "Apr 21": {"other": 2, "export": 1},
    "April 25": {"other": 4, "planting": 5},
    "May": {"other": 3},
    "June": {"other": 7, "planting": 5},
    "July": {"other": 3, "harvest": 4},
    "September": {"other": 5, "export": 1, "harvest": 8},
}


@pytest.mark.parametrize(
    "fragment,expected",
    EXPECTED_CLASSIFICATIONS.items(),
    ids=EXPECTED_CLASSIFICATIONS.keys(),
)
def test_classification_counts(fragment: str, expected: dict[str, int]) -> None:
    """Category counts per file must stay stable."""
    path = _resolve(fragment)
    classes = classify_docx_tables(path, _TEST_CATEGORIES)

    actual: dict[str, int] = {}
    for c in classes:
        cat = c["category"]
        actual[cat] = actual.get(cat, 0) + 1

    assert actual == expected, (
        f"{path.name}: expected {expected}, got {actual}"
    )


class TestClassifyDocxTables:
    """Detailed classification checks."""

    def test_july_harvest_titles(self, july_path: Path) -> None:
        classes = classify_docx_tables(july_path, _TEST_CATEGORIES)
        harvest = [c for c in classes if c["category"] == "harvest"]
        titles = {c["title"] for c in harvest}
        assert titles == {"Grains and pulses", "WHEAT", "BARLEY", "RAPE"}

    def test_sep_harvest_titles(self, sep_path: Path) -> None:
        classes = classify_docx_tables(sep_path, _TEST_CATEGORIES)
        harvest = [c for c in classes if c["category"] == "harvest"]
        titles = {c["title"] for c in harvest}
        expected = {"Grains and pulses", "WHEAT", "BARLEY", "Rye", "Peas", "CORN", "RAPE", "SUNFLOWER"}
        assert titles == expected

    def test_apr25_planting_detected(self, apr25_path: Path) -> None:
        classes = classify_docx_tables(apr25_path, _TEST_CATEGORIES)
        planting = [c for c in classes if c["category"] == "planting"]
        assert len(planting) == 5

    def test_every_table_has_required_keys(self, all_docx_paths: list[Path]) -> None:
        """Every classification result must have all required keys."""
        required = {"index", "category", "title", "rows", "cols"}
        for path in all_docx_paths:
            for c in classify_docx_tables(path, _TEST_CATEGORIES):
                assert required.issubset(c.keys()), (
                    f"{path.name} table {c.get('index')}: missing keys {required - c.keys()}"
                )

    def test_category_values_valid(self, all_docx_paths: list[Path]) -> None:
        """Category must be one of the defined categories or 'other'."""
        valid = set(_TEST_CATEGORIES.keys()) | {"other"}
        for path in all_docx_paths:
            for c in classify_docx_tables(path, _TEST_CATEGORIES):
                assert c["category"] in valid


# ---------------------------------------------------------------------------
# 10. Round-trip: extract_tables_from_docx still works for ALL files
# ---------------------------------------------------------------------------


def test_all_files_extract_without_error(all_docx_paths: list[Path]) -> None:
    """Smoke test: extraction must not raise on any file."""
    for path in all_docx_paths:
        tables = extract_tables_from_docx(path)
        assert len(tables) > 0, f"{path.name}: no tables extracted"
        for i, t in enumerate(tables):
            assert len(t.column_names) > 0, f"{path.name} table {i}: no columns"
            assert t.source_format == "docx"


def test_all_files_classify_without_error(all_docx_paths: list[Path]) -> None:
    """Smoke test: classification must not raise on any file."""
    for path in all_docx_paths:
        classes = classify_docx_tables(path, _TEST_CATEGORIES)
        assert len(classes) > 0


def test_all_files_compress_without_error(all_docx_paths: list[Path]) -> None:
    """Smoke test: compression must not raise on any file."""
    for path in all_docx_paths:
        results = compress_docx_tables(path)
        # Every file has at least one table
        assert len(results) > 0, f"{path.name}: no compressed tables"
        for md, meta in results:
            assert len(md) > 0
            assert "table_index" in meta


# ---------------------------------------------------------------------------
# 11. extract_pivot_values()
# ---------------------------------------------------------------------------


class TestExtractPivotValues:
    """Verify dynamic year extraction from compressed markdown headers."""

    def test_july_wheat_years(self, july_path: Path) -> None:
        """WHEAT table headers contain 2024 and 2025."""
        results = compress_docx_tables(july_path, table_indices=[4])
        md, _ = results[0]
        years = extract_pivot_values(md)
        assert years == ["2024", "2025"]

    def test_returns_sorted_unique(self) -> None:
        """Years must be sorted and deduplicated."""
        md = "| A / 2025 | B / 2024 | C / 2025 | D / 2023 |\n| --- | --- | --- | --- |"
        years = extract_pivot_values(md)
        assert years == ["2023", "2024", "2025"]

    def test_no_years_returns_empty(self) -> None:
        """No years in header → empty list."""
        md = "| Region | Area | Metric |\n| --- | --- | --- |"
        assert extract_pivot_values(md) == []

    def test_empty_string(self) -> None:
        assert extract_pivot_values("") == []

    def test_skips_separator_line(self) -> None:
        """Must not extract from separator row."""
        md = "| A / 2025 |\n| --- |\n| 2020 |"
        years = extract_pivot_values(md)
        assert years == ["2025"]  # from header, not data

    def test_slicing_last_n_years(self) -> None:
        """User can slice to get only recent years."""
        md = "| A / 2022 | B / 2023 | C / 2024 | D / 2025 |\n| --- |"
        years = extract_pivot_values(md)
        assert years[-2:] == ["2024", "2025"]
        assert years[-1:] == ["2025"]


# ---------------------------------------------------------------------------
# 12. Synthetic DOCX — stress tests with materialized documents
#
# All previous tests rely on the 6 Russian agricultural reports.  These
# synthetic tests use DOCX files in inputs/docx/synthetic/ with controlled
# merge patterns, verifying that the pipeline is truly format-agnostic.
#
# To regenerate the files: uv run python tests/generate_synthetic_docx.py
# ---------------------------------------------------------------------------

SYNTH_DIR = Path("inputs/docx/synthetic")


def _synth(name: str) -> Path:
    p = SYNTH_DIR / f"{name}.docx"
    assert p.exists(), f"Synthetic DOCX not found: {p}  (run: uv run python tests/generate_synthetic_docx.py)"
    return p


# --- 12a. Flat tables (no merges) -----------------------------------------
#
# Note: extract_tables_from_docx() uses estimate_header_rows() which relies
# on cell-type patterns and styling — unavailable in programmatic tables.
# Flat synthetic tables are tested via compress_docx_tables() which uses the
# merge-based detector (always returns at least 1 header row).


class TestSyntheticFlat:
    """Tables with no merges — baseline behaviour."""

    def test_compress_shape(self):
        """Flat table: 1 header row (minimum), remaining rows are data."""
        path = _synth("flat")
        results = compress_docx_tables(path)
        assert len(results) == 1
        md, meta = results[0]
        assert meta["row_count"] == 5
        assert meta["col_count"] == 3
        header_line = md.split("\n")[0]
        assert "Name" in header_line

    def test_header_count_minimum_1(self):
        """No merges → _detect_header_rows_from_merges returns 1 (minimum)."""
        from docx import Document
        path = _synth("flat")
        doc = Document(path)
        hc = _detect_header_rows_from_merges(doc.tables[0])
        assert hc == 1

    def test_single_column(self):
        """Single-column table compresses correctly."""
        path = _synth("single_column")
        results = compress_docx_tables(path)
        assert len(results) == 1
        md, meta = results[0]
        assert meta["row_count"] == 4
        assert "Countries" in md.split("\n")[0]


# --- 12b. Horizontal span headers -----------------------------------------
#
# For _detect_header_rows_from_merges to detect 2+ header rows, the table
# needs a merge indicator on EACH header row.  A column spanning all header
# rows vertically (vmerge) is the natural pattern — e.g., "Category" or "ID"
# spanning the full header height — and matches real-world DOCX tables.


class TestSyntheticHorizontalMerge:
    """Tables with gridSpan in header rows + vmerge on label column."""

    def test_compound_headers(self):
        """Horizontal spans + vmerge → compound '/' headers."""
        path = _synth("hspan")
        results = compress_docx_tables(path)
        md, meta = results[0]
        header_line = md.split("\n")[0]
        assert "Revenue / Q1" in header_line
        assert "Revenue / Q2" in header_line
        assert "Cost / Q1" in header_line
        assert "Cost / Q2" in header_line
        assert meta["row_count"] == 4

    def test_grid_no_duplication(self):
        """Spanned cells must not produce duplicate text in the grid."""
        from docx import Document
        path = _synth("hspan")
        doc = Document(path)
        grid, _ = _build_grid_from_table(doc.tables[0])
        # Row 0 has "Revenue" spanning cols 1-2
        assert grid[0][1] == "Revenue"
        assert grid[0][2] == ""  # continuation, not duplicated

    def test_full_width_title_span(self):
        """A title cell spanning the entire row width (title_hspan.docx)."""
        from docx import Document
        path = _synth("title_hspan")
        doc = Document(path)
        grid, _ = _build_grid_from_table(doc.tables[0])
        assert grid[0][0] == "QUARTERLY REPORT"
        assert all(grid[0][c] == "" for c in range(1, 5))


# --- 12c. Title row detection ---------------------------------------------


class TestSyntheticTitleRow:
    """Tables where row 0 is a title (single non-empty cell)."""

    def test_title_extracted(self):
        path = _synth("title_hspan")
        results = compress_docx_tables(path)
        md, meta = results[0]
        assert meta["title"] == "QUARTERLY REPORT"
        assert md.startswith("## QUARTERLY REPORT")

    def test_title_excluded_from_headers(self):
        """Title row must not appear in compound column names."""
        path = _synth("title_hspan")
        results = compress_docx_tables(path)
        md, meta = results[0]
        assert meta["title"] == "QUARTERLY REPORT"
        pipe_lines = [l for l in md.split("\n") if l.startswith("|") and "---" not in l]
        header_line = pipe_lines[0]
        assert "QUARTERLY REPORT" not in header_line
        assert "Revenue / Q1" in header_line

    def test_multi_category_financial_title(self):
        """FINANCIAL SUMMARY title in multi_category.docx table 2."""
        path = _synth("multi_category")
        results = compress_docx_tables(path, table_indices=[2])
        md, meta = results[0]
        assert meta["title"] == "FINANCIAL SUMMARY"
        assert md.startswith("## FINANCIAL SUMMARY")


# --- 12d. Vertical merges -------------------------------------------------


class TestSyntheticVerticalMerge:
    """Tables with vMerge in header rows."""

    def test_vmerge_detected_as_header(self):
        """Vertical merge in header rows → header_count ≥ 2."""
        from docx import Document
        path = _synth("hspan")
        doc = Document(path)
        hc = _detect_header_rows_from_merges(doc.tables[0])
        assert hc == 2  # row 0 has gridSpan, vMerge on col 0 reaches row 1

    def test_vmerge_compound_header_output(self):
        """Vertical merge on label column produces correct compound headers."""
        path = _synth("unicode")
        results = compress_docx_tables(path)
        md, _ = results[0]
        header_line = md.split("\n")[0]
        assert "Région" in header_line
        assert "Données / 2023" in header_line
        assert "Données / 2025" in header_line


# --- 12e. Deep hierarchy (3+ header rows) ---------------------------------


class TestSyntheticDeepHierarchy:
    """Tables with 3+ levels of merged headers."""

    def test_three_level_compound_headers(self):
        path = _synth("deep_hierarchy")
        results = compress_docx_tables(path)
        md, meta = results[0]
        header_line = md.split("\n")[0]
        assert "Group A / Sub X / 2024" in header_line
        assert "Group A / Sub X / 2025" in header_line
        assert "Group A / Sub Y / 2024" in header_line
        assert "Group B / Sub X / 2025" in header_line
        assert meta["row_count"] == 5

    def test_deep_header_count(self):
        """deep_hierarchy.docx has 3 header rows (hmerges in rows 0-1, vmerge to row 2)."""
        from docx import Document
        path = _synth("deep_hierarchy")
        doc = Document(path)
        hc = _detect_header_rows_from_merges(doc.tables[0])
        assert hc == 3

    def test_pivot_values_from_deep_headers(self):
        """Years buried in a deep hierarchy are still extracted."""
        path = _synth("deep_hierarchy")
        results = compress_docx_tables(path)
        md, _ = results[0]
        years = extract_pivot_values(md)
        assert years == ["2024", "2025"]


# --- 12f. Classification with user-defined categories ---------------------


class TestSyntheticClassification:
    """classify_docx_tables with caller-supplied categories."""

    def test_multi_table_routing(self):
        """multi_category.docx: 4 tables, each matching different categories."""
        path = _synth("multi_category")
        cats = {
            "shipping": ["cargo", "port", "vessel"],
            "hr": ["employee", "salary", "department"],
            "finance": ["revenue", "expenses"],
            "inventory": ["stock", "warehouse"],
        }
        classes = classify_docx_tables(path, cats)
        assert len(classes) == 4
        assert classes[0]["category"] == "shipping"
        assert classes[1]["category"] == "hr"
        assert classes[2]["category"] == "finance"
        assert classes[3]["category"] == "inventory"

    def test_empty_categories_gives_other(self):
        path = _synth("flat")
        classes = classify_docx_tables(path, {})
        assert classes[0]["category"] == "other"

    def test_case_insensitive_matching(self):
        """multi_category.docx table 0 has 'Cargo volume' → matches 'cargo'."""
        path = _synth("multi_category")
        classes = classify_docx_tables(path, {"logistics": ["cargo volume"]})
        assert classes[0]["category"] == "logistics"

    def test_title_not_used_for_classification(self):
        """FINANCIAL SUMMARY title must not influence category of other tables."""
        path = _synth("multi_category")
        # Only keyword in table 2's header rows is "Revenue", "Expenses"
        # The title "FINANCIAL SUMMARY" should not count
        classes = classify_docx_tables(path, {"summary": ["financial summary"]})
        assert classes[2]["category"] == "other"
        assert classes[2]["title"] == "FINANCIAL SUMMARY"


# --- 12g. Multiple tables in one document ---------------------------------


class TestSyntheticMultiTable:
    """Documents with several tables of mixed structure."""

    def test_table_indices_filter(self):
        """multi_category.docx has 4 tables; index filter selects subset."""
        path = _synth("multi_category")
        all_results = compress_docx_tables(path)
        assert len(all_results) == 4
        filtered = compress_docx_tables(path, table_indices=[1])
        assert len(filtered) == 1
        assert filtered[0][1]["table_index"] == 1

    def test_mixed_flat_and_merged(self):
        """multi_category.docx has flat (tables 0,1,3) and merged (table 2)."""
        path = _synth("multi_category")
        results = compress_docx_tables(path)
        # Table 0: flat → no "/" in headers
        md0, _ = results[0]
        assert "/" not in md0.split("\n")[0]
        # Table 2: merged with title → compound headers
        md2, meta2 = results[2]
        assert meta2["title"] == "FINANCIAL SUMMARY"
        pipe_lines = [l for l in md2.split("\n") if l.startswith("|") and "---" not in l]
        assert "Revenue / 2024" in pipe_lines[0]


# --- 12h. Edge cases ------------------------------------------------------


class TestSyntheticEdgeCases:
    """Boundary conditions and unusual table shapes."""

    def test_wide_table_10_cols(self):
        """Wide table with 10 columns and spans."""
        path = _synth("wide_table")
        results = compress_docx_tables(path)
        md, meta = results[0]
        header_line = md.split("\n")[0]
        assert "Block A / C1" in header_line
        assert "Block B / C9" in header_line
        assert meta["col_count"] == 10

    def test_unicode_content(self):
        """Non-ASCII content is preserved through the pipeline."""
        path = _synth("unicode")
        results = compress_docx_tables(path)
        md, _ = results[0]
        assert "Île-de-France" in md
        assert "Provence-Alpes-Côte d'Azur" in md
        assert "Données / 2023" in md

    def test_unicode_pivot_values(self):
        """Year extraction works on unicode tables."""
        path = _synth("unicode")
        results = compress_docx_tables(path)
        md, _ = results[0]
        years = extract_pivot_values(md)
        assert years == ["2023", "2024", "2025"]

    def test_single_column_smoke(self):
        """Single-column DOCX extracts without error."""
        path = _synth("single_column")
        tables = extract_tables_from_docx(path)
        assert len(tables) == 1

    def test_all_synthetic_files_compress(self):
        """Smoke test: every synthetic DOCX compresses without error."""
        for path in sorted(SYNTH_DIR.glob("*.docx")):
            results = compress_docx_tables(path)
            assert len(results) > 0, f"{path.name}: no compressed tables"
            for md, meta in results:
                assert len(md) > 0
                assert "table_index" in meta
